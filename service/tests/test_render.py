"""Tests for app.render — the enriched-markdown → branded .docx pipeline.

What runs everywhere (pure Python, no binaries):
  - diagram extraction, caption resolution, node-count heuristics
  - the mermaid.live pako encoding (locked to the byte output of the real
    mermaid-live-editor, and self-consistent via a decode round-trip)
  - frontmatter split + display-only heading cleanup (strip IDs, promote levels)
  - the markdown transform (inline figures, single Caption style) + ```card markers + appendix
  - the editorial post-process: cover (full-bleed hero), tables, stat cards (python-docx; importorskip)

What skips gracefully when the binary is absent (documented in README):
  - actual mmdc / d2 rasterization  → requires mmdc, d2 (+ rsvg-convert)
  - pandoc → .docx                  → requires pandoc
"""

import base64
import json
import re
import shutil
import struct
import zipfile
import zlib
from pathlib import Path

import pytest

from app import render

FIXTURE_DIR = Path(__file__).parent / "fixtures" / "enriched"


@pytest.fixture
def no_binaries(monkeypatch):
    """Force the renderer-absent path regardless of the host so degrade behavior is
    deterministic."""
    monkeypatch.setattr(render, "_which", lambda _b: None)


def _write_png(path, width, height):
    """A structurally valid PNG with the given IHDR dimensions (enough for _png_size /
    python-docx; pixel data is a stub)."""
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)

    def chunk(typ, body):
        return struct.pack(">I", len(body)) + typ + body + struct.pack(">I", zlib.crc32(typ + body) & 0xFFFFFFFF)

    idat = zlib.compress(b"\x00" * (width * 3 + 1))
    with open(path, "wb") as fh:
        fh.write(sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


# --------------------------------------------------------------------------- #
# 1. Diagram extraction
# --------------------------------------------------------------------------- #

def test_extract_diagrams_picks_only_mermaid_and_d2():
    md = (
        "# Title\n\nIntro.\n\n"
        "```mermaid\nflowchart TD\n  A --> B\n```\n\n"
        "Some prose.\n\n"
        "```python\nprint('not a diagram')\n```\n\n"
        "```d2\na -> b\n```\n"
    )
    blocks = render.extract_diagrams(md)
    assert [b.lang for b in blocks] == ["mermaid", "d2"]
    assert "flowchart TD" in blocks[0].code
    assert "a -> b" in blocks[1].code


def test_extract_ignores_unclosed_fence_safely():
    md = "```mermaid\nflowchart TD\n  A --> B\n"  # never closed
    blocks = render.extract_diagrams(md)
    assert len(blocks) == 1
    assert "A --> B" in blocks[0].code


def test_normal_code_fences_pass_through_text():
    md = "```js\nconst x = 1;\n```\n"
    tokens = list(render._iter_blocks(md))
    assert all(k == "text" for k, _ in tokens)
    assert "const x = 1;" in tokens[0][1]


# --------------------------------------------------------------------------- #
# 2. mermaid.live pako encoding (the must-verify item)
# --------------------------------------------------------------------------- #

# Byte-for-byte output of the REAL mermaid-live-editor serde (pako.deflate level 9 +
# js-base64 url-safe) for the state below — captured from a cross-check against the
# editor's own pako + js-base64, and confirmed to round-trip through its deserializer.
_CANONICAL_CODE = "graph TD\n  A[Start] --> B{Decision}\n  B -->|yes| C[Do thing]\n  B -->|no| D[Stop]"
_CANONICAL_URL = (
    "https://mermaid.live/edit#pako:"
    "eNpFjjsOgzAQBa-y2hou4CJSiG9AOkKxsjfYUvyRWRcIuHuwUqSd0dO8HU2yjAqXQtnBU78iwH0a"
    "hYrM0Pc3GHbNxq8-xbO5ocFj4_WAx6QTiPNxmf8mpgP0tU95xg4Dl0DeotpRHIcWsvym-hE8O6Qq"
    "adyiQSWlcoc1WxLWnq4z4QfPL48xNf4"
)


def test_mermaid_live_url_matches_real_editor_bytes():
    assert render.mermaid_live_url(_CANONICAL_CODE) == _CANONICAL_URL


def test_mermaid_live_url_roundtrips_like_the_editor():
    """Decode exactly as mermaid-live-editor does: url-safe b64 (repad) → zlib inflate
    → JSON.parse. Must recover the original state."""
    code = "sequenceDiagram\n  Alice->>Bob: Hi\n  Bob-->>Alice: Hello"
    url = render.mermaid_live_url(code)
    payload = url.split("#pako:", 1)[1]
    payload += "=" * (-len(payload) % 4)
    state = json.loads(zlib.decompress(base64.urlsafe_b64decode(payload)))
    assert state["code"] == code
    assert state["mermaid"] == {"theme": "default"}


def test_mermaid_live_url_is_url_safe_and_unpadded():
    payload = render.mermaid_live_url("graph TD\n A-->B").split("#pako:", 1)[1]
    assert "+" not in payload and "/" not in payload and "=" not in payload


# --------------------------------------------------------------------------- #
# 3. Caption resolution
# --------------------------------------------------------------------------- #

def test_caption_from_fence_attribute(tmp_path):
    md = '```mermaid caption="Component architecture"\nflowchart TD\n A-->B\n```\n'
    _, figs = render.transform(md, doc_stem="d", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"))
    assert figs[0].caption == "Component architecture"


def test_caption_from_in_code_directive_is_stripped(tmp_path):
    cap, code = render._strip_caption_directive("mermaid", "%% caption: X\nflowchart TD\n A-->B")
    assert cap == "X" and "caption" not in code
    md = "```mermaid\n%% caption: From the source\nflowchart TD\n A-->B\n```\n"
    _, figs = render.transform(md, doc_stem="d", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"))
    assert figs[0].caption == "From the source"
    assert "%% caption" not in figs[0].code


def test_caption_falls_back_to_preceding_heading(tmp_path):
    md = "## Data model\n\n```d2\na -> b\n```\n"
    _, figs = render.transform(md, doc_stem="d", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"))
    assert figs[0].caption == "Data model"


def test_caption_uses_single_caption_style_never_bold(tmp_path, no_binaries):
    md = "## Arch\n\n```mermaid\nflowchart TD\n A-->B\n```\n"
    out, _ = render.transform(md, doc_stem="d", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"))
    assert 'custom-style="Caption"' in out      # one uniform Caption style
    assert "**Figure" not in out                # never hand-bolded
    assert "Figure 1 — Arch" in out


# --------------------------------------------------------------------------- #
# 4. Node-count heuristics
# --------------------------------------------------------------------------- #

def test_count_flowchart_nodes():
    code = "flowchart TD\n A[Start] --> B{Q}\n B -->|y| C[Do]\n B -->|n| D[Stop]"
    assert render.count_nodes("mermaid", code) == 4


def test_count_erdiagram_entities():
    code = (
        "erDiagram\n"
        "  CUSTOMER ||--o{ ORDER : places\n"
        "  ORDER ||--|{ LINE_ITEM : contains\n"
        "  CUSTOMER {\n    string name\n  }\n"
    )
    assert render.count_nodes("mermaid", code) == 3


def test_count_sequence_participants():
    code = "sequenceDiagram\n participant A\n participant B\n participant C\n A->>B: x"
    assert render.count_nodes("mermaid", code) == 3


def test_count_d2_nodes():
    code = "a -> b\nb -> c\nd: A box\n"
    assert render.count_nodes("d2", code) == 4  # a, b, c, d


def test_large_diagram_flagged_not_readable_in_transform(tmp_path):
    edges = "\n".join(f"  N{i} --> N{i + 1}" for i in range(20))  # 21 nodes > budget
    md = f"```mermaid\nflowchart TD\n{edges}\n```\n"
    out, figs = render.transform(md, doc_stem="big", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"))
    assert figs[0].node_count > render.NODE_BUDGET
    assert figs[0].not_readable
    assert figs[0].view_url and figs[0].view_url.startswith("https://mermaid.live/edit#pako:")
    assert "View larger" in out


# --------------------------------------------------------------------------- #
# 5. Frontmatter + display-only heading cleanup
# --------------------------------------------------------------------------- #

def test_split_frontmatter_parses_and_drops_status():
    md = ('---\ntitle: Tech Arch\nsubtitle: "Solution Design"\nproject: Acme\n'
          'version: 0.3\ndate: 2026-06-26\nstatus: pre-freeze\n---\n# Doc 3 · Title\n\nBody.\n')
    meta, body = render._split_frontmatter(md)
    assert meta == {"title": "Tech Arch", "subtitle": "Solution Design",
                    "project": "Acme", "version": "0.3", "date": "2026-06-26"}
    assert "status" not in meta
    assert body.startswith("# Doc 3 · Title")


def test_split_frontmatter_absent_is_noop():
    md = "# Doc 3 · Title\n\nBody.\n"
    meta, body = render._split_frontmatter(md)
    assert meta == {} and body == md


def test_transform_headings_strips_ids_and_promotes():
    md = ("# Doc 3 · Technical Architecture\n\n## F3.1 · Solution overview\n\ntext\n\n"
          "### F3.P · Deep dive\n\n## §0 · Introduction\n")
    out, title, _ = render._transform_headings(md)
    assert title == "Technical Architecture"   # `# Doc N · …` captured + cleaned
    assert "# Doc 3" not in out                 # title removed from the body
    assert "# Solution overview" in out         # ## → #, F3.1 stripped
    assert "## Deep dive" in out                # ### → ##, F3.P stripped
    assert "# Introduction" in out              # §0 stripped, promoted
    assert "F3.1" not in out and "§0" not in out and "F3.P" not in out


def test_transform_headings_leaves_fenced_code_alone():
    md = "## F1.1 · Title\n\n```d2\n## not a heading\na -> b\n```\n"
    out, _, _ = render._transform_headings(md)
    assert "## not a heading" in out   # untouched inside the fence
    assert "# Title" in out            # the real heading stripped + promoted


def test_transform_strips_inline_xref_tokens_from_prose():
    """The hard rule: no internal F-section/§ IDs visible in body prose. Parenthetical
    refs drop (prose names the section); bare refs become the mapped section name."""
    md = (
        "## F3.2 · Component architecture\n\n"
        "## F3.4 · Deterministic logic boundary\n\n"
        "## F3.6 · Integration specifications\n\n"
        "## §0 · Introduction\n\n"
        "the deterministic boundary (F3.4) then gates the flow.\n\n"
        "see the F3.6 gap and again see the F3.6 gap.\n\n"
        "specified under the AI/ML approach (F3.3) and Monitoring & drift (F3.12).\n\n"
        'The "CieTrade Sync" component (F3.2) and the pipeline.\n\n'
        "the same human-in-the-loop boundary drawn in F3.4.\n\n"
        "terms are defined in §0 of this doc.\n"
    )
    out, _, _ = render._transform_headings(md)
    body = "\n".join(ln for ln in out.splitlines() if not ln.startswith("#"))
    # zero internal ID tokens in the body prose
    assert not re.search(r"F\d+\.[0-9A-Za-z]+", body)
    assert "§" not in body
    # clean results
    assert "the deterministic boundary then gates the flow." in body          # paren dropped
    assert "see the Integration specifications gap and again see the Integration specifications gap." in body
    assert "specified under the AI/ML approach and Monitoring & drift." in body  # parens dropped
    assert 'The "CieTrade Sync" component and the pipeline.' in body
    assert "drawn in Deterministic logic boundary." in body                    # bare → name
    assert "terms are defined in Introduction of this doc." in body            # §0 → name
    # no spacing/paren litter
    assert " )" not in body and "( " not in body and "()" not in body and "  " not in body


def test_inline_strips_dimension_ids():
    """Dimension IDs ([A-Z]{3}-\\d{2}, RAID-…) are dropped from prose — bare and
    parenthesized — while real values like AES-256 / TLS 1.2 stay intact."""
    md = (
        "## F3.4 · Deterministic logic boundary\n\n"
        "Owner ENG-04 maintains it; the bar is DAT-07 (SEC-01).\n\n"
        "Specified under the approach (ENG-03, SCO-09), pending Assumption RAID-A07.\n\n"
        "Encryption uses AES-256 and TLS 1.2; remittance is EDI 820 over SFTP.\n"
    )
    out, _, _ = render._transform_headings(md)
    assert not re.search(r"(?<![\w-])[A-Z]{3}-\d{2}(?!\d)", out)  # no dimension IDs
    assert "RAID-" not in out                                     # no RAID IDs
    assert "Owner maintains it; the bar is." in out              # bare + paren dropped, tidied
    assert "Specified under the approach, pending Assumption." in out
    assert "AES-256" in out and "TLS 1.2" in out and "EDI 820" in out  # real values intact


def test_strips_source_provenance_spans():
    """*Source: …* spans are removed entirely — standalone, trailing a paragraph, and
    across soft-wrapped lines."""
    md = (
        "## F3.4 · Title\n\n"
        "The roles are defined here. *Source: instance inventory ENG-04; decision log\n"
        "SEC-01, DAT-08.*\n\n"
        "Next paragraph stays.\n\n"
        "*Source: decision log ENG-03, DAT-06.*\n"
    )
    out, _, _ = render._transform_headings(md)
    assert "Source:" not in out
    assert "The roles are defined here." in out      # trailing span removed cleanly
    assert "Next paragraph stays." in out


def test_transform_headings_extracts_opening_note():
    """The leading abstract (prose before the first section heading) is lifted out for the
    Pull-Quote page; inline IDs in it are stripped too."""
    md = ("This document sets out the architecture and the rules in F3.4 in full.\n\n"
          "## F3.4 · Deterministic logic boundary\n\nBody text.\n")
    out, title, note = render._transform_headings(md)
    assert note == ["This document sets out the architecture and the rules in Deterministic logic boundary in full."]
    assert "This document sets out" not in out          # lifted out of the body
    assert "# Deterministic logic boundary" in out       # section promoted to H1


def test_opening_note_follows_doc_title_but_not_section():
    md = "# Doc 3 · Tech Arch\n\nLeading abstract paragraph.\n\n## F3.1 · Overview\n\nBody.\n"
    out, title, note = render._transform_headings(md)
    assert title == "Tech Arch"
    assert note == ["Leading abstract paragraph."]
    assert "Leading abstract paragraph." not in out


def test_no_opening_note_when_body_starts_with_heading():
    md = "# Doc 3 · Tech Arch\n\n## F3.1 · Overview\n\nBody.\n"
    _, _, note = render._transform_headings(md)
    assert note == []


# --------------------------------------------------------------------------- #
# 6. Transform: inline figures, fit-to-page, appendix (degraded — forced)
# --------------------------------------------------------------------------- #

def test_transform_degrades_without_renderers(tmp_path, no_binaries):
    md = ("## Architecture\n\n```mermaid\nflowchart TD\n A-->B\n```\n\n"
          "## Schema\n\n```d2\nx -> y\n```\n")
    warnings: list[str] = []
    out, figs = render.transform(
        md, doc_stem="doc", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"), warnings=warnings,
    )
    assert len(figs) == 2
    assert all(f.png is None for f in figs)            # nothing rasterized
    assert "Figure 1 — Architecture" in out and "Figure 2 — Schema" in out
    assert 'custom-style="Caption"' in out and "**Figure" not in out
    assert "renderer unavailable" in out               # source fence kept visible
    assert "# Figure sources" in out                   # self-contained appendix (H1)
    assert out.count("```mermaid") >= 1 and out.count("```d2") >= 1
    assert warnings                                     # degradation reported, not silent


def test_appendix_uses_caption_style_and_h1():
    figs = [
        render.Figure(number=1, lang="mermaid", caption="One", code="flowchart TD\n A-->B"),
        render.Figure(number=2, lang="d2", caption="Two", code="a -> b"),
    ]
    apx = render._appendix(figs)
    assert "# Figure sources" in apx                    # one TOC-worthy section heading
    assert apx.count('custom-style="Caption"') == 2     # entry labels reuse Caption style
    assert "Figure 1 — One" in apx and "Figure 2 — Two" in apx
    assert "### Figure" not in apx                       # no per-figure sub-headings


def test_no_landscape_artifacts_remain():
    """The landscape post-process and its flags are fully removed."""
    assert not hasattr(render, "apply_landscape")
    assert not hasattr(render.PageGeometry, "landscape_content_in")
    assert "wide" not in {f.name for f in __import__("dataclasses").fields(render.Figure)}


def test_wide_image_clamped_to_portrait_width(tmp_path, monkeypatch):
    """A figure wider than the page content is scaled to the content WIDTH (inline,
    no landscape) and flagged not_readable for a view-larger link."""
    monkeypatch.setattr(render, "_which", lambda b: "/usr/bin/" + b)

    def fake_mmdc(code, out_png, scale=render.SCALE):
        _write_png(out_png, width=int(12 * scale * render.DPI), height=int(4 * scale * render.DPI))
        return render._png_size(out_png)

    monkeypatch.setattr(render, "render_mermaid_png", fake_mmdc)
    out, figs = render.transform(
        "```mermaid\nflowchart LR\n A-->B\n```\n", doc_stem="w",
        figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"),
    )
    assert figs[0].not_readable is True
    assert f"width={render._fmt_in(render.DEFAULT_GEOMETRY.portrait_content_in)}in" in out  # 6.5in


def test_tall_image_clamped_to_page_height(tmp_path, monkeypatch):
    """A very tall figure is scaled so its HEIGHT fits the page (keeps the height-clamp
    even though landscape is gone)."""
    monkeypatch.setattr(render, "_which", lambda b: "/usr/bin/" + b)

    def fake_mmdc(code, out_png, scale=render.SCALE):
        _write_png(out_png, width=int(3 * scale * render.DPI), height=int(20 * scale * render.DPI))
        return render._png_size(out_png)

    monkeypatch.setattr(render, "render_mermaid_png", fake_mmdc)
    out, _ = render.transform(
        "```mermaid\nflowchart TD\n A-->B\n```\n", doc_stem="t",
        figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"),
    )
    w = float(re.search(r"width=([\d.]+)in", out).group(1))
    expected = 3 * (render.DEFAULT_GEOMETRY.portrait_content_h_in / 20)  # height-bound fit
    assert abs(w - expected) < 0.05 and w < 3.0


# --------------------------------------------------------------------------- #
# 7. PNG header reader
# --------------------------------------------------------------------------- #

def test_png_size_reads_ihdr(tmp_path):
    p = tmp_path / "x.png"
    _write_png(str(p), 321, 123)
    assert render._png_size(str(p)) == (321, 123)


# --------------------------------------------------------------------------- #
# 8. Cover page (needs python-docx)
# --------------------------------------------------------------------------- #

def _texts(path):
    import docx
    return [p.text for p in docx.Document(str(path)).paragraphs]


def test_add_cover_page_from_frontmatter(tmp_path):
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_heading("Section One", level=1)
    document.add_paragraph("Body content.")
    path = tmp_path / "d.docx"
    document.save(str(path))

    meta = {"title": "Technical Architecture", "subtitle": "Solution Design",
            "project": "Acme", "version": "0.3", "date": "2026-06-26"}
    assert render.add_cover_page(str(path), meta=meta) is True

    texts = _texts(path)
    blob = "\n".join(texts)
    # subtitle renders as the mono eyebrow (uppercased)
    for v in ("Technical Architecture", "SOLUTION DESIGN", "Acme", "Version 0.3", "2026-06-26"):
        assert v in blob
    assert texts.index("Technical Architecture") < texts.index("Body content.")  # cover first
    xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode("utf-8", "replace")
    assert 'w:val="nextPage"' in xml  # cover is its own section ⇒ stands alone


def test_add_cover_page_falls_back_to_doc_title(tmp_path):
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_paragraph("Body.")
    path = tmp_path / "d.docx"
    document.save(str(path))
    assert render.add_cover_page(str(path), meta={}, title_fallback="Fallback Title") is True
    assert "Fallback Title" in "\n".join(_texts(path))


def test_add_cover_page_omits_status(tmp_path):
    pytest.importorskip("docx")
    import docx
    document = docx.Document()
    document.add_paragraph("Body.")
    path = tmp_path / "d.docx"
    document.save(str(path))
    render.add_cover_page(str(path), meta={"title": "T", "status": "pre-freeze"})
    blob = "\n".join(_texts(path))
    assert "T" in blob and "pre-freeze" not in blob


def test_add_cover_page_noop_when_nothing_to_show(tmp_path):
    pytest.importorskip("docx")
    import docx
    document = docx.Document()
    document.add_paragraph("Body.")
    path = tmp_path / "d.docx"
    document.save(str(path))
    assert render.add_cover_page(str(path), meta={}, title_fallback=None) is False


def test_cover_with_hero_image_is_fullbleed(tmp_path):
    docx = pytest.importorskip("docx")
    img = tmp_path / "cover.png"
    _write_png(str(img), 850, 1100)  # full-page-aspect hero stand-in
    document = docx.Document()
    document.add_heading("Body", level=1)
    document.add_paragraph("Content.")
    path = tmp_path / "d.docx"
    document.save(str(path))

    meta = {"title": "Tech Arch", "subtitle": "Solution Design",
            "project": "Acme", "version": "0.3", "date": "2026-06-27"}
    assert render.add_cover_page(str(path), meta=meta, cover_image_path=str(img)) is True

    z = zipfile.ZipFile(str(path))
    xml = z.read("word/document.xml").decode("utf-8", "replace")
    assert "<wp:anchor" in xml and 'behindDoc="1"' in xml          # full-bleed background
    assert any(n.startswith("word/media/") for n in z.namelist())  # hero embedded
    blob = "\n".join(_texts(path))
    assert "Tech Arch" in blob and "SOLUTION DESIGN" in blob       # title + mono eyebrow
    assert 'w:val="nextPage"' in xml                               # cover is its own section


def test_cover_logo_centered_above_title(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    hero = tmp_path / "cover.png"; _write_png(str(hero), 850, 1100)
    logo = tmp_path / "logo.png"; _write_png(str(logo), 2048, 700)
    document = docx.Document(); document.add_heading("Body", level=1); document.add_paragraph("Content.")
    path = tmp_path / "d.docx"; document.save(str(path))

    assert render.add_cover_page(str(path), meta={"title": "Tech Arch"}, cover_image_path=str(hero),
                                 logo_path=str(logo)) is True
    z = zipfile.ZipFile(str(path))
    xml = z.read("word/document.xml").decode("utf-8", "replace")
    assert xml.count("<wp:anchor") == 1                    # only the hero floats; logo is inline
    assert 'name="cover-hero"' in xml and 'name="cover-logo"' not in xml
    assert len([n for n in z.namelist() if n.startswith("word/media/")]) >= 2  # hero + logo embedded
    # the logo (inline picture) is centered and sits just above the title
    doc = docx.Document(str(path))
    paras = doc.paragraphs
    logo_idx = next(i for i, p in enumerate(paras) if p._p.findall(".//" + qn("wp:inline")))
    title_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "Tech Arch")
    assert logo_idx < title_idx
    assert paras[logo_idx].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_cover_note_is_vertically_centered_pullquote_page(tmp_path):
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_heading("Section One", level=1)
    document.add_paragraph("Body content.")
    path = tmp_path / "d.docx"; document.save(str(path))

    assert render.add_cover_page(str(path), meta={"title": "T"},
                                 note=["The opening abstract."]) is True
    texts = _texts(path)
    assert "The opening abstract." in texts
    assert texts.index("The opening abstract.") < texts.index("Section One")  # note before body
    xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode("utf-8", "replace")
    assert xml.count('w:val="nextPage"') >= 2   # cover + note are their own sections
    assert 'w:val="center"' in xml              # the note section is vertically centered


def test_section_starts_on_new_page_after_toc(tmp_path):
    """With a TOC (a w:sdt as the first block), a page break is inserted after it so
    Section 1 begins on a fresh page."""
    docx = pytest.importorskip("docx")
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = docx.Document()
    heading = document.add_heading("Section One", level=1)
    # Emulate pandoc --toc: an sdt as the very first body block.
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    content.append(OxmlElement("w:p"))
    sdt.append(content)
    document.element.body.insert(0, sdt)
    path = tmp_path / "d.docx"; document.save(str(path))

    assert render.add_cover_page(str(path), meta={"title": "T"}) is True
    body = docx.Document(str(path)).element.body
    kids = list(body)
    sdt_idx = next(i for i, el in enumerate(kids) if el.tag == qn("w:sdt"))
    after = kids[sdt_idx + 1]
    assert after.tag == qn("w:p") and after.findall(".//" + qn("w:br")), "page break expected after TOC"


# --------------------------------------------------------------------------- #
# 9. Stat / callout cards (```card block)
# --------------------------------------------------------------------------- #

def test_parse_card_handles_colon_in_value():
    assert render._parse_card("value: 1\nlabel: X\nnote: a: b") == {"value": "1", "label": "X", "note": "a: b"}


def test_card_block_is_intercepted_not_a_diagram(tmp_path):
    md = "# T\n\n```card\nvalue: 99.5%\nlabel: AVAILABILITY\nnote: monthly\n```\n"
    assert render.extract_diagrams(md) == []                       # a card is not a figure
    out, figs = render.transform(
        md, doc_stem="c", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"),
    )
    assert figs == []
    m = re.search(r"\[\[KEEL-CARD:([A-Za-z0-9_\-=]+)\]\]", out)
    assert m
    fields = json.loads(base64.urlsafe_b64decode(m.group(1)).decode("utf-8"))
    assert fields == {"value": "99.5%", "label": "AVAILABILITY", "note": "monthly"}


def test_render_cards_builds_grey_card_box(tmp_path):
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_paragraph("Intro.")
    document.add_paragraph(render._card_marker("value: 99.5%\nlabel: AVAILABILITY\nnote: Monthly SLA.").strip())
    document.add_paragraph("After.")
    path = tmp_path / "c.docx"
    document.save(str(path))

    assert render.render_cards(str(path)) == 1
    reopened = docx.Document(str(path))
    assert len(reopened.tables) == 1
    cell_texts = [p.text for p in reopened.tables[0].cell(0, 0).paragraphs]
    assert "99.5%" in cell_texts and "AVAILABILITY" in cell_texts and "Monthly SLA." in cell_texts
    xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode("utf-8", "replace")
    assert 'w:fill="F2EEE4"' in xml                 # grey card box
    assert "454BC4" in xml and "Oranienbaum" in xml  # brand-purple Oranienbaum value
    assert "KEEL-CARD" not in xml                    # marker consumed


# --------------------------------------------------------------------------- #
# 10. Editorial tables (post-process)
# --------------------------------------------------------------------------- #

def test_style_tables_editorial(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = docx.Document()
    t = document.add_table(rows=4, cols=2)
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Metric", "Target"
    for ri, (a, b) in enumerate([("Latency", "≤ 200 ms"), ("Availability", "99.9%"),
                                 ("Throughput", "1,100")], start=1):
        t.rows[ri].cells[0].text, t.rows[ri].cells[1].text = a, b
    path = tmp_path / "t.docx"
    document.save(str(path))

    assert render.style_tables(str(path)) == 1
    xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode("utf-8", "replace")
    assert 'w:fill="15171A"' in xml                                  # dark header (per-cell)
    assert "FFFFFF" in xml and "IBM Plex Mono SemiBold" in xml       # white mono header text
    assert 'w:fill="FBF9F4"' in xml and 'w:fill="F2EEE4"' in xml     # zebra body rows
    assert 'w:color="E6E1D6"' in xml                                 # hairline borders
    assert 'w:after="320"' in xml                                    # space after the table

    reopened = docx.Document(str(path))
    tbl = reopened.tables[0]
    # numeric Target column right-aligned; text Metric column not
    assert tbl.rows[1].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert tbl.rows[1].cells[0].paragraphs[0].alignment != WD_ALIGN_PARAGRAPH.RIGHT


def test_style_tables_drops_all_empty_column(tmp_path):
    """A column whose body is entirely empty (e.g. a provenance column whose only content
    was dimension IDs, now blanked) is dropped."""
    docx = pytest.importorskip("docx")
    document = docx.Document()
    t = document.add_table(rows=3, cols=3)
    t.rows[0].cells[0].text, t.rows[0].cells[1].text, t.rows[0].cells[2].text = "Metric", "Target", "Source"
    t.rows[1].cells[0].text, t.rows[1].cells[1].text = "Latency", "99.9%"   # "Source" body left empty
    t.rows[2].cells[0].text, t.rows[2].cells[1].text = "Uptime", "99.5%"
    path = tmp_path / "t.docx"
    document.save(str(path))

    render.style_tables(str(path))
    tbl = docx.Document(str(path)).tables[0]
    assert len(tbl.columns) == 2                                  # empty "Source" column dropped
    assert [c.text for c in tbl.rows[0].cells] == ["Metric", "Target"]


def test_table_cell_with_only_dimension_id_is_blanked():
    md = ("## F3.7 · NFRs\n\n| Metric | Owner |\n|---|---|\n"
          "| Latency | ENG-07 |\n| Uptime | DAT-07 |\n")
    out, _, _ = render._transform_headings(md)
    assert "ENG-07" not in out and "DAT-07" not in out           # ref-only cells blanked
    assert "Latency" in out and "Uptime" in out                  # the data column survives


# --------------------------------------------------------------------------- #
# 11. Orchestration (degraded: no pandoc → clean rendered.md, no docx)
# --------------------------------------------------------------------------- #

def test_render_markdown_file_strips_ids_and_frontmatter(tmp_path, no_binaries):
    src = tmp_path / "3-tech.md"
    src.write_text(
        "---\ntitle: Tech Arch\nproject: Acme\n---\n# Doc 3 · Technical Architecture\n\n"
        "## F3.1 · Overview\n\nText.\n\n## F3.2 · Components\n\nMore.\n",
        encoding="utf-8",
    )
    res = render.render_markdown_file(str(src), str(tmp_path / "out"))
    rendered = Path(res.rendered_md).read_text(encoding="utf-8")
    assert "F3.1" not in rendered and "F3.2" not in rendered  # IDs stripped for display
    assert "# Overview" in rendered and "# Components" in rendered  # ## → # (H1)
    assert "# Doc 3" not in rendered                          # title → cover, removed
    assert "title: Tech Arch" not in rendered                 # frontmatter consumed
    assert res.docx is None                                   # pandoc absent → degrade
    # The on-disk source is never modified (gate reads it by ID).
    assert "## F3.1 · Overview" in src.read_text(encoding="utf-8")


def test_render_markdown_file_zero_inline_id_leaks(tmp_path, no_binaries):
    """End-to-end: the rendered.md body carries 0 F-section/§ tokens (the hard rule)."""
    src = tmp_path / "3-tech.md"
    src.write_text(
        "# Doc 3 · Technical Architecture\n\n"
        "## F3.4 · Deterministic logic boundary\n\n"
        "## F3.6 · Integration specifications\n\n"
        "The deterministic boundary (F3.4) gates the flow; see the F3.6 gap. "
        "It reuses the boundary drawn in F3.4.\n",
        encoding="utf-8",
    )
    res = render.render_markdown_file(str(src), str(tmp_path / "out"))
    body = "\n".join(
        ln for ln in Path(res.rendered_md).read_text(encoding="utf-8").splitlines()
        if not ln.lstrip().startswith(("*Source", "```"))
    )
    assert not re.search(r"F\d+\.[0-9]", body) and not re.search(r"§[0-9]", body)


def test_enriched_fixture_transform(tmp_path, no_binaries):
    """The committed enriched fixture (with frontmatter) exercises caption mechanisms,
    the node-budget heuristic, the d2 path, and the uniform Caption style."""
    meta, body = render._split_frontmatter((FIXTURE_DIR / "3-technical-architecture.md").read_text(encoding="utf-8"))
    assert meta.get("title") and meta.get("project")  # fixture carries frontmatter
    out, figs = render.transform(body, doc_stem="f3", figures_dir=str(tmp_path / "f"), attachments_dir=str(tmp_path / "a"))
    assert [f.lang for f in figs] == ["mermaid", "mermaid", "mermaid", "d2"]
    assert figs[0].caption == "System-in-context"          # fence attribute
    assert figs[3].caption == "Rules engine internals"     # in-code directive (stripped)
    assert "caption" not in figs[3].code
    erd = figs[2]
    assert erd.node_count > render.NODE_BUDGET and erd.not_readable and erd.view_url
    assert "# Figure sources" in out
    assert out.count('custom-style="Caption"') == 8        # 4 inline + 4 appendix labels


# --------------------------------------------------------------------------- #
# 10. Real-binary smoke tests (skip when the tool is absent)
# --------------------------------------------------------------------------- #

@pytest.mark.skipif(not shutil.which(render.MMDC_BIN), reason="mmdc (mermaid-cli) not installed")
def test_mmdc_renders_png(tmp_path):
    out = tmp_path / "m.png"
    w, h = render.render_mermaid_png("flowchart TD\n A-->B\n B-->C", str(out))
    assert out.exists() and w > 0 and h > 0


@pytest.mark.skipif(
    not (shutil.which(render.D2_BIN) and shutil.which(render.RSVG_BIN)),
    reason="d2 and/or rsvg-convert not installed",
)
def test_d2_renders_png_and_svg(tmp_path):
    png, svg = tmp_path / "d.png", tmp_path / "d.svg"
    render.render_d2_svg("a -> b -> c", str(svg))
    w, h = render.render_d2_png("a -> b -> c", str(png))  # svg → rsvg-convert
    assert svg.exists() and png.exists() and w > 0 and h > 0


@pytest.mark.skipif(not shutil.which(render.PANDOC_BIN), reason="pandoc not installed")
def test_full_pipeline_cover_and_clean_headings(tmp_path):
    src = tmp_path / "deliverables"
    src.mkdir()
    (src / "3-tech.md").write_text(
        "---\ntitle: Technical Architecture\nproject: Acme\nversion: 0.3\n---\n"
        "# Doc 3 · Technical Architecture\n\n## F3.1 · Overview\n\n"
        "```mermaid\nflowchart TD\n A-->B\n```\n",
        encoding="utf-8",
    )
    out = tmp_path / "out"
    results = render.render_directory(str(src), str(out))
    assert len(results) == 1 and results[0].docx and (out / "3-tech.docx").exists()
    xml = zipfile.ZipFile(str(out / "3-tech.docx")).read("word/document.xml").decode("utf-8", "replace")
    assert "Technical Architecture" in xml          # cover title
    assert 'w:type="page"' in xml                    # cover page break
    assert "F3.1" not in xml                          # IDs stripped for display
    assert 'w:orient="landscape"' not in xml         # no landscape sections
