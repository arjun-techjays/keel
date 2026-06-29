"""Render enriched markdown deliverables (with diagrams) into branded Word .docx.

This module is **standalone and environment-agnostic**: it imports nothing from the
service (no Supabase, no config), so it runs both as a CLI

    python -m app.render <deliverables_dir> --reference-doc <ref.docx> --out <dir>

and as a library imported by ``checkout_core`` for the service render route.

Pipeline (per deliverable, in order):
  1. Split off YAML frontmatter (title/subtitle/project/version/date) for the cover.
  2. Clean headings for DISPLAY ONLY (source markdown unchanged on disk): strip the
     internal ``F{doc}.{n}`` / ``§`` / ``Doc N`` prefixes and promote every heading one
     level (## → #, ### → ##) so sections become Heading 1 and Word auto-numbers them.
  3. Extract fenced ```mermaid / ```d2 blocks; render each to PNG at 3x with a
     LOCAL/PRIVATE renderer (mmdc for mermaid; d2 → SVG → rsvg-convert for d2).
  4. Replace each block with the image (inline, at content width; text flows after) +
     a single uniformly-styled ``Figure N — <caption>`` caption. Tall/over-wide figures
     are scaled to fit the page; over-budget/over-wide ones get a mermaid.live "view
     larger" link (mermaid) or a high-res SVG attachment link (dense d2).
  5. Convert ```card blocks to inert markers (rendered to stat cards in step 8).
  6. Append a "Figure sources" appendix so every doc is self-contained (durability).
  7. Convert to .docx via pandoc + the branding reference doc, with a TOC.
  8. Editorial post-process with python-docx (the look pandoc + the reference doc can't
     do): restyle every table (dark mono header, zebra rows, padding, space-after,
     numeric right-align), turn card markers into stat cards, and prepend a designed
     full-page cover (hero image + title block) from the frontmatter.

Diagrams render INLINE at content width and text flows after them — no page rotation,
no forced one-figure-per-page (decomposition keeps figures small; the render contract).

HARD PRIVACY RULE (non-negotiable): diagram source is rendered ONLY by local binaries
(``mmdc`` for mermaid, the ``d2`` binary for d2). It is NEVER sent to a public service
(kroki.io, mermaid.ink, …). The mermaid.live "view larger" link is a *string we build
locally and embed* — the source rides in the URL ``#fragment``, which browsers never
transmit to a server, and this module never fetches it. There are **zero outbound
network calls anywhere in this module** (note the import list below — no httpx/requests).
"""

from __future__ import annotations

import argparse
import base64
import contextlib
import copy
import dataclasses
import json
import os
import re
import shutil
import struct
import subprocess
import sys
import tempfile
import zlib
from pathlib import Path

# ---------------------------------------------------------------------------
# Constants & geometry
# ---------------------------------------------------------------------------

SCALE = 3                 # raster diagrams at 3x for crisp print/zoom (render contract)
NODE_BUDGET = 15          # complexity budget — over this, flag "not readable" (render contract)
DPI = 96                  # px-per-inch pandoc assumes for PNGs without DPI metadata
_EPS = 0.05               # inch tolerance so a figure exactly at content width isn't "wide"

DIAGRAM_LANGS = ("mermaid", "d2")
# Fenced blocks the render intercepts (others pass through to pandoc untouched).
_INTERCEPT = {"mermaid": "diagram", "d2": "diagram", "card": "card"}

# Environment overrides for the local binaries (privacy: local only, never a URL).
MMDC_BIN = os.environ.get("KEEL_MMDC_BIN", "mmdc")
D2_BIN = os.environ.get("KEEL_D2_BIN", "d2")
RSVG_BIN = os.environ.get("KEEL_RSVG_BIN", "rsvg-convert")
PANDOC_BIN = os.environ.get("KEEL_PANDOC_BIN", "pandoc")
MERMAID_LIVE_THEME = os.environ.get("KEEL_MERMAID_LIVE_THEME", "default")

# Editorial palette + fonts (design-system.md §2–3), applied by the python-docx
# post-process — the effects pandoc + the reference doc can't do alone. Values mirror
# the approved type specimen (assets/branding/specimens/type-specimen.docx).
COL_INK = "1C1C1E"            # headings + table data
COL_BODY = "34343A"          # body prose (softened)
COL_MUTED = "5A5A5E"         # captions, labels, footer
COL_BRAND = "454BC4"         # card values, section numbers, links
COL_PAPER = "FBF9F4"         # warm paper tint (table zebra A)
COL_ZEBRA = "F2EEE4"         # card fill + table zebra B
COL_TABLE_HEADER = "15171A"  # near-black table header fill (white mono text)
COL_WHITE = "FFFFFF"
COL_HAIRLINE = "E6E1D6"      # table borders / rules
DISPLAY_FONT = "Oranienbaum"
BODY_FONT = "Roboto"
MONO_FONT = "IBM Plex Mono"
MONO_SEMIBOLD = "IBM Plex Mono SemiBold"


@dataclasses.dataclass(frozen=True)
class PageGeometry:
    """Page size & margins — the content box figures are scaled to fit (portrait only;
    figures render inline, never rotated). Defaults: US Letter, 1in margins."""

    page_w_in: float = 8.5
    page_h_in: float = 11.0
    margin_in: float = 1.0

    @property
    def portrait_content_in(self) -> float:
        return self.page_w_in - 2 * self.margin_in

    @property
    def portrait_content_h_in(self) -> float:
        return self.page_h_in - 2 * self.margin_in


DEFAULT_GEOMETRY = PageGeometry()


class RenderError(RuntimeError):
    """A required step failed (pandoc error, renderer error when strict)."""


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclasses.dataclass
class DiagramBlock:
    lang: str           # "mermaid" | "d2"
    code: str           # diagram source (caption directive stripped)
    info: str = ""      # raw fence info string (e.g. 'mermaid caption="…"')


@dataclasses.dataclass
class Figure:
    number: int
    lang: str
    caption: str
    code: str
    png: str | None = None        # path to rendered PNG (None if renderer unavailable)
    svg: str | None = None        # path to high-res SVG (d2 only)
    width_in: float | None = None  # natural width of the 1x render, in inches
    height_in: float | None = None  # natural height of the 1x render, in inches
    node_count: int = 0
    not_readable: bool = False     # over node budget OR wider than the page content
    view_url: str | None = None    # mermaid.live link (mermaid, when not_readable)


@dataclasses.dataclass
class RenderResult:
    src: str
    rendered_md: str
    docx: str | None
    figures: list[Figure]
    warnings: list[str] = dataclasses.field(default_factory=list)
    tables: int = 0
    cards: int = 0


# ---------------------------------------------------------------------------
# 1. Diagram extraction
# ---------------------------------------------------------------------------

_FENCE_OPEN = re.compile(r"^(\s{0,3})(`{3,}|~{3,})[ \t]*(.*?)[ \t]*$")


def _iter_blocks(md: str):
    """Yield ('text', str), ('diagram', {...}) and ('card', {...}) tokens in order.

    A robust line scanner (not a single regex) so the exact source span of each
    intercepted fence is preserved and ordinary code fences pass through untouched.
    """
    lines = md.splitlines()
    i, n = 0, len(lines)
    text_buf: list[str] = []
    while i < n:
        m = _FENCE_OPEN.match(lines[i])
        info = m.group(3) if m else ""
        lang = info.split()[0].lower() if (m and info.strip()) else ""
        kind = _INTERCEPT.get(lang) if m else None
        if kind:
            fchar = m.group(2)[0]
            flen = len(m.group(2))
            close_re = re.compile(r"^\s{0,3}" + re.escape(fchar) + r"{" + str(flen) + r",}\s*$")
            j = i + 1
            code_lines: list[str] = []
            closed = False
            while j < n:
                if close_re.match(lines[j]):
                    closed = True
                    break
                code_lines.append(lines[j])
                j += 1
            if text_buf:
                yield ("text", "\n".join(text_buf))
                text_buf = []
            yield (kind, {"lang": lang, "code": "\n".join(code_lines), "info": info})
            i = (j + 1) if closed else j
        else:
            text_buf.append(lines[i])
            i += 1
    if text_buf:
        yield ("text", "\n".join(text_buf))


def extract_diagrams(md: str) -> list[DiagramBlock]:
    """All mermaid/d2 diagram blocks in document order (caption directive stripped)."""
    out: list[DiagramBlock] = []
    for kind, payload in _iter_blocks(md):
        if kind == "diagram":
            _, code = _strip_caption_directive(payload["lang"], payload["code"])
            out.append(DiagramBlock(lang=payload["lang"], code=code, info=payload["info"]))
    return out


# ---------------------------------------------------------------------------
# 2. Captions
# ---------------------------------------------------------------------------

_CAP_ATTR = re.compile(r"caption\s*=\s*(?:\"([^\"]*)\"|'([^']*)')")
_CAP_HTML = re.compile(r"<!--\s*caption:\s*(.+?)\s*-->", re.I)
_HEADING = re.compile(r"^\s{0,3}#{1,6}\s+(.*?)\s*#*\s*$")
# A caption directive *inside* the diagram: mermaid uses %% comments, d2 uses #.
_CAP_IN_CODE = {
    "mermaid": re.compile(r"^\s*%%\s*caption:\s*(.+?)\s*$", re.I),
    "d2": re.compile(r"^\s*#\s*caption:\s*(.+?)\s*$", re.I),
}


def _strip_caption_directive(lang: str, code: str) -> tuple[str | None, str]:
    """Pull a ``caption:`` comment line out of the diagram source. Returns
    (caption_or_None, code_without_that_line)."""
    rx = _CAP_IN_CODE.get(lang)
    if not rx:
        return None, code
    caption = None
    kept: list[str] = []
    for ln in code.splitlines():
        m = rx.match(ln)
        if m and caption is None:
            caption = m.group(1).strip()
            continue
        kept.append(ln)
    return caption, "\n".join(kept)


def _resolve_caption(info: str, in_code: str | None, preceding_text: str, number: int) -> str:
    """Caption precedence: fence attr → in-diagram directive → preceding HTML
    comment → nearest preceding heading → generic ``Figure N``."""
    m = _CAP_ATTR.search(info or "")
    if m:
        return (m.group(1) or m.group(2)).strip()
    if in_code:
        return in_code
    if preceding_text:
        hm = _CAP_HTML.search(preceding_text)
        if hm:
            return hm.group(1).strip()
        for ln in reversed(preceding_text.splitlines()):
            hd = _HEADING.match(ln)
            if hd and hd.group(1).strip():
                return hd.group(1).strip()
    return f"Figure {number}"


def _caption_block(text: str) -> str:
    """One figure caption / appendix label, as a pandoc fenced div bound to the single
    'Caption' paragraph style — uniform for every figure, never hand-bolded (contract)."""
    return f'::: {{custom-style="Caption"}}\n{text}\n:::'


# A ```card block carries its data to the docx post-process as an inert marker
# paragraph: pandoc passes the literal text through (no markdown specials in url-safe
# base64), and render_cards() turns each marker into a styled 1-cell card table.
_CARD_MARKER_RE = re.compile(r"^\[\[KEEL-CARD:([A-Za-z0-9_\-=]+)\]\]$")


def _parse_card(code: str) -> dict:
    """Parse a ```card block's ``key: value`` lines (value · label · note · …)."""
    fields: dict[str, str] = {}
    for ln in code.splitlines():
        key, sep, val = ln.partition(":")
        if sep and key.strip():
            fields[key.strip().lower()] = val.strip()
    return fields


def _card_marker(code: str) -> str:
    payload = base64.urlsafe_b64encode(
        json.dumps(_parse_card(code), separators=(",", ":")).encode("utf-8")
    ).decode("ascii")
    return f"\n\n[[KEEL-CARD:{payload}]]\n\n"


# ---------------------------------------------------------------------------
# 3. mermaid.live "view larger" link (verified byte-identical to the live editor)
# ---------------------------------------------------------------------------


def mermaid_live_url(code: str, theme: str = MERMAID_LIVE_THEME) -> str:
    """Build a ``https://mermaid.live/edit#pako:<payload>`` URL.

    The payload exactly matches the mermaid-live-editor "pako" serde:
      JSON.stringify(state)  →  pako.deflate(level 9) [= zlib stream, NOT raw]  →
      js-base64 url-safe (no padding).  Verified byte-identical against the editor's
      own pako + js-base64 and round-tripped through its deserializer, so the link
      loads the diagram. The source lives in the URL #fragment — never sent to any
      server, and this function performs no network I/O.
    """
    state = {"code": code, "mermaid": {"theme": theme}, "autoSync": True, "updateDiagram": True}
    raw = json.dumps(state, separators=(",", ":")).encode("utf-8")  # no spaces (matches JS)
    payload = base64.urlsafe_b64encode(zlib.compress(raw, 9)).rstrip(b"=").decode("ascii")
    return f"https://mermaid.live/edit#pako:{payload}"


# ---------------------------------------------------------------------------
# 4. Node-count heuristics (best-effort, per diagram family)
# ---------------------------------------------------------------------------

_ARROW_RE = re.compile(r"<?[-.=]{2,}>?|--+|==+|-\.+-|~~+")
_ID_TOKEN = re.compile(r"[A-Za-z_][\w./-]*")


def _mermaid_kind(code: str) -> str:
    for ln in code.splitlines():
        s = ln.strip()
        if not s or s.startswith("%%"):
            continue
        return s.split()[0].lower()
    return ""


def count_mermaid_nodes(code: str) -> int:
    kind = _mermaid_kind(code)
    lines = [ln for ln in code.splitlines() if ln.strip() and not ln.strip().startswith("%%")]
    body = "\n".join(lines[1:]) if lines else ""

    if kind in ("flowchart", "graph"):
        nodes: set[str] = set()
        for tok in re.findall(r"([A-Za-z_]\w*)\s*[\[\(\{]", body):
            nodes.add(tok)
        for ln in body.splitlines():
            parts = _ARROW_RE.split(ln)
            if len(parts) > 1:
                for p in parts:
                    p = re.sub(r"[\[\(\{].*", "", p).strip().strip("|").strip()
                    p = p.split("|")[-1].strip()
                    if p and re.fullmatch(r"[A-Za-z_]\w*", p):
                        nodes.add(p)
        return len(nodes)

    if kind == "erdiagram":
        ents = set(re.findall(r"^\s*([A-Za-z_]\w*)\s*\{", body, re.M))
        for ln in body.splitlines():
            m = re.match(r"\s*([A-Za-z_]\w*)\s+[|}o{<>.\-]+\s+([A-Za-z_]\w*)\s*:", ln)
            if m:
                ents.add(m.group(1))
                ents.add(m.group(2))
        return len(ents)

    if kind in ("sequencediagram",):
        parts = set(re.findall(r"^\s*(?:participant|actor)\s+([A-Za-z_]\w*)", body, re.M))
        if parts:
            return len(parts)
        names: set[str] = set()
        for m in re.finditer(r"([A-Za-z_]\w*)\s*-{1,2}>>?\s*([A-Za-z_]\w*)", body):
            names.add(m.group(1))
            names.add(m.group(2))
        return len(names)

    if kind == "classdiagram":
        classes = set(re.findall(r"\bclass\s+([A-Za-z_]\w*)", body))
        for m in re.finditer(r"([A-Za-z_]\w*)\s*[<|*o\-\.]{2,}[>|*o\-\.]*\s*([A-Za-z_]\w*)", body):
            classes.add(m.group(1))
            classes.add(m.group(2))
        return len(classes)

    if kind in ("statediagram", "statediagram-v2"):
        states: set[str] = set()
        for ln in body.splitlines():
            parts = _ARROW_RE.split(ln)
            if len(parts) > 1:
                for p in parts:
                    p = p.split(":")[0].strip()
                    if p and re.fullmatch(r"[A-Za-z_]\w*", p) and p != "[*]":
                        states.add(p)
        return len(states)

    if kind == "gantt":
        return sum(1 for ln in body.splitlines() if ":" in ln and not ln.strip().lower().startswith(("section", "title", "dateformat", "axisformat", "excludes")))

    # Fallback: meaningful (non-empty, non-comment, non-directive) lines.
    return sum(1 for ln in lines[1:] if not ln.strip().startswith(("title", "%%")))


def count_d2_nodes(code: str) -> int:
    nodes: set[str] = set()
    for raw in code.splitlines():
        ln = raw.split("#", 1)[0].strip()  # strip d2 line comments
        if not ln or ln in ("{", "}"):
            continue
        conn = re.split(r"<->|->|<-|--", ln)
        if len(conn) > 1:
            for end in conn:
                end = end.split(":")[0].strip().strip("{").strip()
                if end and _ID_TOKEN.fullmatch(end.split(".")[0]):
                    nodes.add(end.split(".")[0])
            continue
        m = re.match(r"^([A-Za-z_][\w.-]*)\s*[:{]", ln)
        if m:
            nodes.add(m.group(1).split(".")[0])
    return len(nodes)


def count_nodes(lang: str, code: str) -> int:
    return count_mermaid_nodes(code) if lang == "mermaid" else count_d2_nodes(code)


# ---------------------------------------------------------------------------
# Local renderers (subprocess; privacy: offline binaries only)
# ---------------------------------------------------------------------------


def _which(binary: str) -> str | None:
    return shutil.which(binary) or (binary if os.path.isabs(binary) and os.path.exists(binary) else None)


def renderer_status() -> dict:
    """Which local tools are available — for diagnostics and graceful skipping."""
    return {
        "mmdc": _which(MMDC_BIN),
        "d2": _which(D2_BIN),
        "rsvg-convert": _which(RSVG_BIN),
        "pandoc": _which(PANDOC_BIN),
    }


def _png_size(path: str) -> tuple[int, int]:
    """(width, height) in px from the PNG IHDR header — no Pillow dependency."""
    with open(path, "rb") as fh:
        head = fh.read(26)
    if head[:8] != b"\x89PNG\r\n\x1a\n" or head[12:16] != b"IHDR":
        raise RenderError(f"not a PNG: {path}")
    w, h = struct.unpack(">II", head[16:24])
    return w, h


def _mmdc_puppeteer_config() -> str | None:
    """In a container mmdc runs as root, where Chromium refuses the sandbox; pass
    ``--no-sandbox`` then. Returns a temp config path, or None when not needed."""
    needs = bool(os.environ.get("KEEL_MMDC_NO_SANDBOX")) or (hasattr(os, "geteuid") and os.geteuid() == 0)
    if not needs:
        return None
    fd, path = tempfile.mkstemp(suffix=".json", prefix="mmdc-pptr-")
    with os.fdopen(fd, "w", encoding="utf-8") as fh:
        json.dump({"args": ["--no-sandbox", "--disable-setuid-sandbox"]}, fh)
    return path


def render_mermaid_png(code: str, out_png: str, scale: int = SCALE) -> tuple[int, int]:
    """Render mermaid → PNG at ``scale``x via mmdc (local). Returns (w, h) in px."""
    mmdc = _which(MMDC_BIN)
    if not mmdc:
        raise RenderError("mmdc (mermaid-cli) not found — install @mermaid-js/mermaid-cli")
    pptr = _mmdc_puppeteer_config()
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.mmd")
        Path(src).write_text(code, encoding="utf-8")
        cmd = [mmdc, "-i", src, "-o", out_png, "-s", str(scale), "-b", "white"]
        if pptr:
            cmd += ["-p", pptr]
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        finally:
            if pptr:
                with contextlib.suppress(OSError):
                    os.unlink(pptr)
    if proc.returncode != 0 or not os.path.exists(out_png):
        raise RenderError(f"mmdc failed: {(proc.stderr or proc.stdout).strip()[:500]}")
    return _png_size(out_png)


def render_d2_svg(code: str, out_svg: str) -> None:
    """Render d2 → SVG (native, vector) for the high-res attachment / durability."""
    d2 = _which(D2_BIN)
    if not d2:
        raise RenderError("d2 binary not found — install https://d2lang.com")
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.d2")
        Path(src).write_text(code, encoding="utf-8")
        proc = subprocess.run([d2, src, out_svg], capture_output=True, text=True, timeout=120)
    if proc.returncode != 0 or not os.path.exists(out_svg):
        raise RenderError(f"d2 (svg) failed: {(proc.stderr or proc.stdout).strip()[:500]}")


def rasterize_svg_to_png(svg_path: str, out_png: str, scale: int = SCALE) -> tuple[int, int]:
    """SVG → PNG at ``scale``x via rsvg-convert (librsvg). Returns (w, h) in px.

    Deliberately browserless: this is how d2 figures get their inline raster. d2's
    own ``.png`` export pulls a Playwright Chromium and crashes as root in a
    container; rsvg-convert is a tiny, reliable apt/brew package with no such issue."""
    rsvg = _which(RSVG_BIN)
    if not rsvg:
        raise RenderError("rsvg-convert (librsvg) not found — install librsvg / librsvg2-bin")
    cmd = [rsvg, "-z", str(scale), svg_path, "-o", out_png]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if proc.returncode != 0 or not os.path.exists(out_png):
        raise RenderError(f"rsvg-convert failed: {(proc.stderr or proc.stdout).strip()[:500]}")
    return _png_size(out_png)


def render_d2_png(code: str, out_png: str, scale: int = SCALE) -> tuple[int, int]:
    """Render d2 → PNG at ``scale``x: native d2 SVG, then rasterize with rsvg-convert.
    Returns (w, h) in px."""
    with tempfile.TemporaryDirectory() as tmp:
        svg = os.path.join(tmp, "d.svg")
        render_d2_svg(code, svg)
        return rasterize_svg_to_png(svg, out_png, scale)


# ---------------------------------------------------------------------------
# 3+4+. Markdown transform: blocks → image + caption (+ readability link), appendix
# ---------------------------------------------------------------------------


def _fmt_in(x: float) -> str:
    return f"{x:.2f}".rstrip("0").rstrip(".")


def _figure_markdown(fig: Figure, geometry: PageGeometry, rel_png: str | None, rel_svg: str | None) -> str:
    """The replacement markdown for one diagram: inline image at content width, text
    flows after it, one Caption-styled caption (+ a view-larger link when flagged)."""
    extras: list[str] = []
    if fig.lang == "mermaid" and fig.not_readable and fig.view_url:
        extras.append(f"[View larger ↗]({fig.view_url})")
    if fig.lang == "d2" and fig.not_readable and rel_svg:
        extras.append(f"[Full-resolution SVG ↗]({rel_svg})")
    caption_text = f"Figure {fig.number} — {fig.caption}"
    if extras:
        caption_text += " · " + " · ".join(extras)

    if rel_png:
        # Scale to fit the page content box (portrait), preserving aspect and never
        # upscaling — so a wide or tall figure fits the page without overflow. No
        # rotation, no forced page break: text flows naturally after the figure.
        box_w, box_h = geometry.portrait_content_in, geometry.portrait_content_h_in
        nat_w = fig.width_in or box_w
        nat_h = fig.height_in or nat_w
        fit = min(box_w / nat_w, box_h / nat_h, 1.0)
        img = f"![]({rel_png}){{width={_fmt_in(nat_w * fit)}in}}"
        return f"\n{img}\n\n{_caption_block(caption_text)}\n"

    # Renderer unavailable → keep the source visible (pandoc renders the fence as a
    # code block) and still number/caption it; the appendix + mermaid.live link keep
    # it useful and durable.
    fence = f"```{fig.lang}\n{fig.code}\n```"
    caption_text += " — diagram source (renderer unavailable)"
    return f"\n{fence}\n\n{_caption_block(caption_text)}\n"


def _appendix(figures: list[Figure]) -> str:
    """The self-contained "Figure sources" backstop. Heading is a top-level (H1)
    section; each entry label reuses the single Caption style (consistent with the
    in-body captions), so only the one section heading lands in the TOC."""
    if not figures:
        return ""
    parts = ["", "", "---", "", "# Figure sources", "",
             "*Every diagram's source, appended so this document is self-contained — "
             "re-editable offline, independent of any third-party viewer.*", ""]
    for fig in figures:
        parts.append(_caption_block(f"Figure {fig.number} — {fig.caption}"))
        parts.append("")
        parts.append(f"```{fig.lang}")
        parts.append(fig.code)
        parts.append("```")
        if fig.lang == "mermaid" and fig.view_url:
            parts.append("")
            parts.append(f"[Open in mermaid.live ↗]({fig.view_url})")
        parts.append("")
    return "\n".join(parts)


def transform(
    md: str,
    *,
    doc_stem: str,
    figures_dir: str,
    attachments_dir: str,
    figures_rel: str = "figures",
    attachments_rel: str = "attachments",
    geometry: PageGeometry = DEFAULT_GEOMETRY,
    scale: int = SCALE,
    node_budget: int = NODE_BUDGET,
    strict: bool = False,
    warnings: list[str] | None = None,
) -> tuple[str, list[Figure]]:
    """Replace every diagram block with an image+caption, render PNGs/SVGs locally,
    apply the readability heuristic, and append the Figure-sources appendix.

    Returns (transformed_markdown, figures). When ``strict`` is False and a renderer
    is missing, the diagram degrades to a visible source fence (no exception)."""
    warnings = warnings if warnings is not None else []
    os.makedirs(figures_dir, exist_ok=True)
    os.makedirs(attachments_dir, exist_ok=True)

    figures: list[Figure] = []
    out_parts: list[str] = []
    preceding_text = ""
    number = 0

    for kind, payload in _iter_blocks(md):
        if kind == "text":
            out_parts.append(payload)
            preceding_text = payload
            continue

        if kind == "card":
            out_parts.append(_card_marker(payload["code"]))
            preceding_text = ""
            continue

        number += 1
        lang = payload["lang"]
        in_cap, code = _strip_caption_directive(lang, payload["code"])
        caption = _resolve_caption(payload["info"], in_cap, preceding_text, number)
        fig = Figure(number=number, lang=lang, caption=caption, code=code,
                     node_count=count_nodes(lang, code))

        png_name = f"{doc_stem}-fig-{number}.png"
        svg_name = f"{doc_stem}-fig-{number}.svg"
        png_path = os.path.join(figures_dir, png_name)
        svg_path = os.path.join(attachments_dir, svg_name)
        rel_png = rel_svg = None

        try:
            if lang == "mermaid":
                w_px, h_px = render_mermaid_png(code, png_path, scale)
                fig.png = png_path
                fig.width_in, fig.height_in = w_px / (scale * DPI), h_px / (scale * DPI)
                rel_png = f"{figures_rel}/{png_name}"
            else:  # d2 — native SVG (durable attachment), then rasterize that SVG
                render_d2_svg(code, svg_path)
                fig.svg = svg_path
                rel_svg = f"{attachments_rel}/{svg_name}"
                w_px, h_px = rasterize_svg_to_png(svg_path, png_path, scale)
                fig.png = png_path
                fig.width_in, fig.height_in = w_px / (scale * DPI), h_px / (scale * DPI)
                rel_png = f"{figures_rel}/{png_name}"
        except RenderError as exc:
            if strict:
                raise
            warnings.append(f"figure {number} ({lang}): {exc}")

        # Readability heuristic: over the node budget, or wider than the page content
        # box (it'll be shrunk to fit width and may get small) → offer a view-larger.
        too_wide = fig.width_in is not None and fig.width_in > geometry.portrait_content_in + _EPS
        fig.not_readable = (fig.node_count > node_budget) or too_wide
        if lang == "mermaid" and (fig.not_readable or fig.png is None):
            fig.view_url = mermaid_live_url(code)

        out_parts.append(_figure_markdown(fig, geometry, rel_png, rel_svg))
        figures.append(fig)
        preceding_text = ""

    body = "\n".join(out_parts)
    return body + _appendix(figures), figures


# ---------------------------------------------------------------------------
# 5. pandoc → .docx
# ---------------------------------------------------------------------------


def to_docx(rendered_md_path: str, out_docx_path: str, reference_doc: str | None = None,
            toc: bool = True) -> None:
    """Convert markdown → branded .docx via pandoc. Run with cwd = the markdown's
    directory so relative ``figures/…`` image paths resolve."""
    pandoc = _which(PANDOC_BIN)
    if not pandoc:
        raise RenderError("pandoc not found — install pandoc")
    work = os.path.dirname(os.path.abspath(rendered_md_path))
    cmd = [pandoc, os.path.basename(rendered_md_path)]
    if reference_doc and os.path.exists(reference_doc):
        cmd += ["--reference-doc", os.path.abspath(reference_doc)]
    if toc:
        cmd += ["--toc"]
    cmd += ["-o", os.path.basename(out_docx_path)]
    proc = subprocess.run(cmd, cwd=work, capture_output=True, text=True, timeout=300)
    if proc.returncode != 0 or not os.path.exists(os.path.join(work, os.path.basename(out_docx_path))):
        raise RenderError(f"pandoc failed: {(proc.stderr or proc.stdout).strip()[:500]}")


# ---------------------------------------------------------------------------
# Doc-level transforms: frontmatter + clean/numbered headings (display only)
# ---------------------------------------------------------------------------

# A leading internal section prefix the render strips for display (the IDs stay in the
# source markdown for traceability + the gate, which matches by ID OR title).
_SECTION_PREFIX = re.compile(r"^(?:Doc\s+\d+|F\d[\w.–-]*|§\d+)\s*·\s*")
_SECTION_PREFIX_CAP = re.compile(r"^(Doc\s+\d+|F\d[\w.–-]*|§\d+)\s*·\s*")  # captures the ID
_ATX = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")

# Internal reference tokens that must NEVER show in the client-facing body (display-only
# strip — the source markdown keeps them all for the gate):
#   - F-section IDs (F3.4, F3.12, F3.P, F2.19) + § IDs (§0) — resolve to the section name,
#   - catalog dimension IDs ([A-Z]{3}-\d{2}: ENG-04, DAT-07, SEC-01 …) + RAID-[ADRQ]nn —
#     dropped. The (?!\d) guard keeps real values like AES-256 / TLS 1.2 intact.
_XREF_TOKEN = r"F\d+\.[0-9A-Za-z]+(?:[–-][0-9A-Za-z]+)?|§\d+"          # F-section / §
_DIM_TOKEN = r"[A-Z]{3}-\d{2}(?!\d)|RAID-[ADRQ]\d*"                    # dimension / RAID
_REF_TOKEN = r"(?:" + _XREF_TOKEN + r"|" + _DIM_TOKEN + r")"          # any internal ref
_ID_RE = re.compile(r"(?<![\w.\-])(?:" + _REF_TOKEN + r")(?![\w])")
# A parenthetical made up of ONLY id tokens (+ separators) — "(F3.4)", "(ENG-03, SCO-09)",
# "(F3.3 and F3.12)" — dropped whole, with the space before it.
_PAREN_IDS_RE = re.compile(
    r"[ \t]*\(\s*(?:" + _REF_TOKEN + r")(?:\s*(?:,|;|/|&|·|and)\s*(?:" + _REF_TOKEN + r"))*\s*\)"
)
# A ``*Source: …*`` provenance span (internal traceability, not client-facing) — removed
# wholesale, whether standalone or trailing a paragraph (and across soft-wrapped lines).
_SOURCE_SPAN_RE = re.compile(r"[ \t]*\*Source:[^*]*\*", re.I)


def _strip_section_prefix(text: str) -> str:
    return _SECTION_PREFIX.sub("", text, count=1).strip()


def _build_section_names(md: str) -> dict[str, str]:
    """Map each section's internal ID → its clean human title (from the headings in this
    same doc), so inline ``F3.4`` references can resolve to the section name. Fence-aware."""
    names: dict[str, str] = {}
    fence: str | None = None
    for ln in md.splitlines():
        if fence is not None:
            if re.match(r"^\s{0,3}" + re.escape(fence[0]) + r"{" + str(len(fence)) + r",}\s*$", ln):
                fence = None
            continue
        fm = _FENCE_OPEN.match(ln)
        if fm and fm.group(2):
            fence = fm.group(2)
            continue
        m = _ATX.match(ln)
        if not m:
            continue
        hm = _SECTION_PREFIX_CAP.match(m.group(2))
        if hm:
            tok = hm.group(1).strip()
            name = m.group(2)[hm.end():].strip()
            if name and (tok.startswith("F") or tok.startswith("§")):
                names[tok] = name
    return names


def _strip_inline_ids(line: str, section_names: dict[str, str]) -> str:
    """Strip inline internal reference tokens from one body line / table-cell line (display
    only). Parenthetical-only refs drop (the prose already names the section); bare F/§ refs
    become the mapped section name; bare dimension IDs drop; spacing/punctuation tidied. A
    table cell that was only a dimension ID is left blank. Lines with no token are untouched."""
    if not _ID_RE.search(line):
        return line
    lead = re.match(r"^(\s*)", line).group(1)
    s = line[len(lead):]
    s = _PAREN_IDS_RE.sub("", s)                                    # drop "(F3.4)" / "(ENG-04)"
    s = _ID_RE.sub(lambda m: section_names.get(m.group(0), ""), s)  # bare F/§ → name | dim → drop
    s = re.sub(r"\(\s*\)", "", s)                                   # orphan empty parens
    s = re.sub(r"\s+([.,;:!?)])", r"\1", s)                         # space before punct
    s = re.sub(r"(\()\s+", r"\1", s)                                # space after "("
    s = re.sub(r" {2,}", " ", s).strip()                           # collapse + trim
    return lead + s


def _split_frontmatter(md: str) -> tuple[dict, str]:
    """Pull a leading YAML frontmatter block (``---`` … ``---``) off the top. Returns
    (meta, body). Simple scalar ``key: value`` parsing (no PyYAML dependency); quotes
    stripped; ``status`` is dropped on principle (never shown). No frontmatter ⇒ ({}, md)."""
    lines = md.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, md
    end = None
    for i in range(1, len(lines)):
        if lines[i].strip() in ("---", "..."):
            end = i
            break
    if end is None:
        return {}, md
    meta: dict[str, str] = {}
    for ln in lines[1:end]:
        if not ln.strip() or ln.lstrip().startswith("#"):
            continue
        key, sep, val = ln.partition(":")
        if not sep:
            continue
        key = key.strip().lower()
        val = val.strip().strip('"').strip("'").strip()
        if key and val and key != "status":
            meta[key] = val
    body = "\n".join(lines[end + 1:])
    return meta, body.lstrip("\n")


def _clean_note(lines: list[str], section_names: dict[str, str]) -> list[str]:
    """Tidy the captured opening-note lines into clean paragraphs (blank-line separated),
    inline cross-refs stripped — for the Pull-Quote front page."""
    txt = "\n".join(lines).strip()
    if not txt:
        return []
    paras = []
    for block in re.split(r"\n\s*\n", txt):
        flat = re.sub(r"\s+", " ", block.replace("\n", " ")).strip()
        flat = _strip_inline_ids(flat, section_names)
        if flat:
            paras.append(flat)
    return paras


def _transform_headings(md: str) -> tuple[str, str | None, list[str]]:
    """Display-only cleanup (the on-disk markdown is never modified). Returns
    (body, title, note_paragraphs):
      - capture the leading ``# Doc N · …`` as the document title (for the cover) and
        drop that line from the body,
      - capture the **opening note** — the leading prose before the first SECTION heading
        (the ``# Doc N`` title doesn't count) — for the Pull-Quote front page, and drop it
        from the body,
      - strip the internal ``F{doc}.{n}`` / ``§n`` / ``Doc N`` prefix from every heading,
      - promote every remaining heading one level (## → #, ### → ##) so sections are
        Heading 1 and Word's multilevel list auto-numbers them 1, 1.1, 2 …,
      - remove ``*Source: …*`` provenance spans entirely (internal traceability), and
      - strip inline internal IDs out of body PROSE and table cells — F/§ refs resolve to
        the section name, dimension IDs ([A-Z]{3}-\\d{2}, RAID-…) drop. The hard rule is no
        internal IDs visible anywhere in the body.
    Fence-aware: lines inside ``` / ~~~ code fences are left untouched."""
    md = _SOURCE_SPAN_RE.sub("", md)  # drop *Source: …* provenance (display only)
    section_names = _build_section_names(md)
    out: list[str] = []
    note_lines: list[str] = []
    title: str | None = None
    title_taken = False
    before_first_section = True  # leading region (note); the title heading doesn't end it
    fence: str | None = None     # active fence run (e.g. "```"), or None when outside

    for ln in md.splitlines():
        if fence is not None:
            (note_lines if before_first_section else out).append(ln)
            if re.match(r"^\s{0,3}" + re.escape(fence[0]) + r"{" + str(len(fence)) + r",}\s*$", ln):
                fence = None
            continue
        fm = _FENCE_OPEN.match(ln)
        if fm and fm.group(2):
            fence = fm.group(2)
            (note_lines if before_first_section else out).append(ln)
            continue
        m = _ATX.match(ln)
        if not m:
            if before_first_section:
                note_lines.append(ln)                        # opening note (lifted out)
            else:
                out.append(_strip_inline_ids(ln, section_names))  # body prose / table cells
            continue
        level = len(m.group(1))
        text = _strip_section_prefix(m.group(2))
        if level == 1 and not title_taken:
            title = text or None
            title_taken = True
            continue  # the document title becomes the cover; drop it from the body
        before_first_section = False  # first real section heading ends the note region
        out.append("#" * max(1, level - 1) + " " + text)
    return "\n".join(out), title, _clean_note(note_lines, section_names)


# ---------------------------------------------------------------------------
# Editorial post-process (python-docx — the look pandoc + the reference doc can't do)
# ---------------------------------------------------------------------------


def _set_cell_fill(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_run(run, *, font=None, size_pt=None, color=None, tracking=None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    if font:
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if tracking is not None:  # letter-spacing, in twips
        rPr = run._element.get_or_add_rPr()
        sp = rPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            rPr.append(sp)
        sp.set(qn("w:val"), str(tracking))


def _restyle_cell(cell, *, font, size_pt, color) -> None:
    for p in cell.paragraphs:
        for run in p.runs:
            _set_run(run, font=font, size_pt=size_pt, color=color)


def _set_table_borders(tbl, color: str | None, sz: int = 4) -> None:
    """All borders single + ``color`` (sz in eighths of a point), or all ``none`` when
    ``color`` is None (used for the borderless card)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tblPr = tbl._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        if color is None:
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
        else:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(sz))
            e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _set_cell_margins(tbl, top: int, bottom: int, left: int, right: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tblPr = tbl._tbl.tblPr
    old = tblPr.find(qn("w:tblCellMar"))
    if old is not None:
        tblPr.remove(old)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)


def _table_width_auto(tbl) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")


_NUMERIC_RE = re.compile(r"^[≤≥<>~=±\s]*[\$€£+-]?\d[\d.,\s]*\s*(%|[A-Za-z/µ°]{1,12})?$")


def _is_numeric(text: str) -> bool:
    return bool(_NUMERIC_RE.match(text.strip()))


def _drop_empty_columns(tbl) -> int:
    """Drop any column whose body cells are all empty — e.g. a provenance/ref column whose
    only content was dimension IDs (now blanked by the inline strip). Never nukes the whole
    table. Returns the count dropped."""
    from docx.oxml.ns import qn
    rows = tbl.rows
    if len(rows) < 2:
        return 0
    ncols = len(rows[0].cells)
    if ncols < 2:
        return 0
    drop = []
    for ci in range(ncols):
        body = []
        for row in rows[1:]:
            try:
                body.append(row.cells[ci].text.strip())
            except IndexError:
                pass
        if body and not any(body):       # at least one body row, all empty
            drop.append(ci)
    if not drop or len(drop) >= ncols:   # never drop every column
        return 0
    for ci in sorted(drop, reverse=True):
        for row in rows:
            tcs = row._tr.findall(qn("w:tc"))
            if ci < len(tcs):
                row._tr.remove(tcs[ci])
        grid = tbl._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            cols = grid.findall(qn("w:gridCol"))
            if ci < len(cols):
                grid.remove(cols[ci])
    return len(drop)


def style_tables(docx_path: str) -> int:
    """Editorial restyle of every pandoc table — the dark-header approach that worked in
    the specimen (per-cell fill + explicit white mono text, NOT a table-style firstRow
    band): near-black `#15171A` header with white IBM Plex Mono; zebra body rows (paper /
    grey); hairline borders; generous padding; right-aligned numeric columns; space after
    each table. Returns the count of tables styled."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    tables = list(doc.tables)
    for tbl in tables:
        if not tbl.rows:
            continue
        _drop_empty_columns(tbl)   # drop all-dim-ID (now-blank) columns before restyling
        rows = tbl.rows
        if not rows:
            continue
        _set_table_borders(tbl, COL_HAIRLINE, sz=4)
        _set_cell_margins(tbl, top=130, bottom=130, left=160, right=160)
        for cell in rows[0].cells:            # dark header, white mono
            _set_cell_fill(cell, COL_TABLE_HEADER)
            _restyle_cell(cell, font=MONO_SEMIBOLD, size_pt=8.5, color=COL_WHITE)
        for di, row in enumerate(rows[1:]):   # zebra body, ink mono
            for cell in row.cells:
                _set_cell_fill(cell, COL_PAPER if di % 2 == 0 else COL_ZEBRA)
                _restyle_cell(cell, font=MONO_FONT, size_pt=9, color=COL_INK)
        ncols = len(rows[0].cells)
        for ci in range(ncols):               # right-align all-numeric columns
            vals = []
            for row in rows[1:]:
                try:
                    vals.append(row.cells[ci].text.strip())
                except IndexError:
                    pass
            nonempty = [v for v in vals if v]
            if nonempty and all(_is_numeric(v) for v in nonempty):
                for row in rows:
                    try:
                        for p in row.cells[ci].paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    except IndexError:
                        pass
        spacer = OxmlElement("w:p")           # space after the table (design-system §4)
        pPr = OxmlElement("w:pPr")
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:after"), "320")          # 16pt
        pPr.append(sp)
        spacer.append(pPr)
        tbl._tbl.addnext(spacer)
    doc.save(docx_path)
    return len(tables)


def _build_card_table(doc, fields: dict):
    """A stat/callout card: borderless single-cell grey box — Oranienbaum value in brand
    purple · mono tracked label · Roboto note (mirrors the specimen card)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    _set_table_borders(tbl, None)
    _set_cell_margins(tbl, top=210, bottom=210, left=260, right=260)
    _table_width_auto(tbl)
    cell = tbl.cell(0, 0)
    _set_cell_fill(cell, COL_ZEBRA)
    if fields.get("value"):
        _set_run(cell.paragraphs[0].add_run(str(fields["value"])),
                 font=DISPLAY_FONT, size_pt=32, color=COL_BRAND)
    if fields.get("label"):
        _set_run(cell.add_paragraph().add_run(str(fields["label"])),
                 font=MONO_SEMIBOLD, size_pt=8.5, color=COL_MUTED, tracking=12)
    if fields.get("note"):
        _set_run(cell.add_paragraph().add_run(str(fields["note"])),
                 font=BODY_FONT, size_pt=9.5, color=COL_BODY)
    return tbl


def render_cards(docx_path: str) -> int:
    """Replace every ``[[KEEL-CARD:…]]`` marker paragraph with a styled card table.
    Returns the count of cards rendered."""
    from docx import Document
    from docx.shared import Pt

    doc = Document(docx_path)
    markers = [p for p in doc.paragraphs if _CARD_MARKER_RE.match(p.text.strip())]
    rendered = 0
    for p in markers:
        m = _CARD_MARKER_RE.match(p.text.strip())
        try:
            fields = json.loads(base64.urlsafe_b64decode(m.group(1)).decode("utf-8"))
        except Exception:
            continue
        tbl = _build_card_table(doc, fields)
        p._p.addprevious(tbl._tbl)            # move the card to the marker's spot
        for run in list(p.runs):              # blank the marker → it becomes space-after
            run._element.getparent().remove(run._element)
        p.paragraph_format.space_after = Pt(16)
        rendered += 1
    doc.save(docx_path)
    return rendered


# ---------------------------------------------------------------------------
# Cover page (python-docx — a designed first page from the frontmatter + hero)
# ---------------------------------------------------------------------------


def _add_floating_image(paragraph, image_path: str, *, width_in: float, x_in: float,
                        y_in: float, behind: bool, height_in: float | None = None,
                        name: str = "img", el_id: int = 9000) -> None:
    """Add ``image_path`` as a floating, page-anchored image to ``paragraph`` — full-bleed
    background (behind text) or a positioned overlay (in front). Raises on any failure so
    the caller can fall back."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu

    run = paragraph.add_run()
    kw = {"width": Emu(int(width_in * 914400))}
    if height_in is not None:
        kw["height"] = Emu(int(height_in * 914400))
    run.add_picture(image_path, **kw)
    inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))

    anchor = OxmlElement("wp:anchor")
    for k, v in (("behindDoc", "1" if behind else "0"), ("distT", "0"), ("distB", "0"),
                 ("distL", "0"), ("distR", "0"), ("simplePos", "0"), ("locked", "0"),
                 ("layoutInCell", "1"), ("allowOverlap", "1"),
                 ("relativeHeight", "0" if behind else "5")):
        anchor.set(k, v)
    simple = OxmlElement("wp:simplePos"); simple.set("x", "0"); simple.set("y", "0")
    anchor.append(simple)
    for axis, off_in in (("wp:positionH", x_in), ("wp:positionV", y_in)):
        pos = OxmlElement(axis); pos.set("relativeFrom", "page")
        off = OxmlElement("wp:posOffset"); off.text = str(int(off_in * 914400))
        pos.append(off); anchor.append(pos)
    for tag in ("wp:extent", "wp:effectExtent"):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)
    anchor.append(OxmlElement("wp:wrapNone"))
    docPr = inline.find(qn("wp:docPr"))
    if docPr is not None:                       # unique id/name so multiple anchors coexist
        docPr.set("id", str(el_id))
        docPr.set("name", name)
        anchor.append(docPr)
    for tag in ("wp:cNvGraphicFramePr", "a:graphic"):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)
    inline.getparent().replace(inline, anchor)


_SECTPR_ORDER = ("w:headerReference", "w:footerReference", "w:footnotePr", "w:endnotePr",
                 "w:type", "w:pgSz", "w:pgMar", "w:paperSrc", "w:pgBorders", "w:lnNumType",
                 "w:pgNumType", "w:cols", "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
                 "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid", "w:printerSettings",
                 "w:sectPrChange")


def _sectpr_insert(sectPr, child) -> None:
    """Insert ``child`` into ``sectPr`` keeping OOXML schema child order."""
    from docx.oxml.ns import qn
    order = [qn(t) for t in _SECTPR_ORDER]
    try:
        mine = order.index(child.tag)
    except ValueError:
        sectPr.append(child); return
    for existing in sectPr:
        if existing.tag in order and order.index(existing.tag) > mine:
            existing.addprevious(child); return
    sectPr.append(child)


def _make_section_sectpr(body_sectPr, *, vertical_center: bool = False):
    """A section-break sectPr cloned from the document's geometry/headers that starts the
    NEXT section on a new page; optionally vertically centers this section's page."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    sp = copy.deepcopy(body_sectPr) if body_sectPr is not None else OxmlElement("w:sectPr")
    t = sp.find(qn("w:type"))
    if t is None:
        t = OxmlElement("w:type"); _sectpr_insert(sp, t)
    t.set(qn("w:val"), "nextPage")
    old = sp.find(qn("w:vAlign"))
    if old is not None:
        sp.remove(old)
    if vertical_center:
        va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "center"); _sectpr_insert(sp, va)
    return sp


def _put_sectpr(p_el, sectPr) -> None:
    """Place ``sectPr`` in paragraph ``p_el``'s pPr (ending that section at this paragraph)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr"); p_el.insert(0, pPr)
    old = pPr.find(qn("w:sectPr"))
    if old is not None:
        pPr.remove(old)
    pPr.append(sectPr)


def add_cover_page(docx_path: str, *, meta: dict, title_fallback: str | None = None,
                   cover_image_path: str | None = None, logo_path: str | None = None,
                   note: list[str] | None = None) -> bool:
    """Build the front matter: a designed full-page **cover** (hero background + the logo
    centred just above a centred title block), then the **opening note** lifted onto its
    own **vertically-centred Pull-Quote page**, then the **TOC and Section 1 each on a new
    page**. Returns True if a cover was added.

    Order produced: Cover (p1) → Note (Pull-Quote page, vertically centred) → TOC (own
    page) → Section 1. Cover/note are their own sections (section breaks); the note section
    carries ``w:vAlign="center"``. Title/eyebrow/note use the template's named styles."""
    from docx import Document  # lazy: only the cover step needs python-docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    title = (meta.get("title") or title_fallback or "").strip()
    eyebrow = (meta.get("subtitle") or "").strip()
    has_image = bool(cover_image_path and os.path.exists(cover_image_path))
    has_logo = bool(logo_path and os.path.exists(logo_path))
    note = note or []
    if not (title or eyebrow or any(meta.get(k) for k in ("project", "version", "date"))
            or has_image or note):
        return False

    doc = Document(docx_path)
    body = doc.element.body
    if len(body) == 0:
        return False
    anchor = body[0]                       # the TOC sdt (with --toc) or first content heading
    body_sectPr = body.find(qn("w:sectPr"))  # clone its geometry/headers for our sections
    created: list = []
    fullbleed = False

    def _para(text="", *, style=None, size=None, color=None, font=None, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=0, space_after=6, tracking=None):
        p = doc.add_paragraph()
        p.alignment = align
        if style:
            with contextlib.suppress(KeyError):
                p.style = doc.styles[style]
        if text:
            _set_run(p.add_run(text), font=font, size_pt=size, color=color, tracking=tracking)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        created.append(p._p)
        return p

    # --- Cover (page 1): full-bleed hero + centred logo + centred title block ---
    if has_image:
        bg = doc.add_paragraph()
        try:
            _add_floating_image(bg, cover_image_path, width_in=DEFAULT_GEOMETRY.page_w_in,
                                height_in=DEFAULT_GEOMETRY.page_h_in, x_in=0, y_in=0,
                                behind=True, name="cover-hero", el_id=9001)
            fullbleed = True
        except Exception:  # fallback: large top image, title beneath
            bg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            with contextlib.suppress(Exception):
                bg.add_run().add_picture(cover_image_path, height=Inches(4.2))
        created.append(bg._p)

    if fullbleed:
        for _ in range(2):  # drop the title block into the hero's upper (light) field
            _para(space_after=8)

    if has_logo:  # logo centred, just above the title block
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        with contextlib.suppress(Exception):
            lp.add_run().add_picture(logo_path, width=Inches(1.7))
        lp.paragraph_format.space_after = Pt(14)
        created.append(lp._p)

    if eyebrow:
        _para(eyebrow.upper(), font=MONO_SEMIBOLD, size=9, color=COL_MUTED, tracking=20, space_after=10)
    if title:
        _para(title, style="Title", color=COL_INK, space_after=10)
    if meta.get("project"):
        _para(str(meta["project"]), font=BODY_FONT, size=13, color=COL_BODY, space_after=4)
    if meta.get("version"):
        _para(f"Version {meta['version']}", font=MONO_FONT, size=9.5, color=COL_MUTED, space_after=2)
    if meta.get("date"):
        _para(str(meta["date"]), font=MONO_FONT, size=9.5, color=COL_MUTED, space_after=2)

    # End the cover as its own section (→ the note/TOC starts on a fresh page).
    cover_close = doc.add_paragraph()
    _put_sectpr(cover_close._p, _make_section_sectpr(body_sectPr))
    created.append(cover_close._p)

    # --- Opening note: its own vertically-centred Pull-Quote page ---
    if note:
        note_ps = []
        for para in note:
            p = doc.add_paragraph()
            with contextlib.suppress(KeyError):
                p.style = doc.styles["Pull Quote"]
            p.add_run(para)
            created.append(p._p); note_ps.append(p._p)
        # the note's own section, vertically centred (sectPr on its last paragraph)
        _put_sectpr(note_ps[-1], _make_section_sectpr(body_sectPr, vertical_center=True))

    for el in created:  # relocate the front matter to the very front, in order
        anchor.addprevious(el)

    # --- Section 1 starts on a new page after the TOC ---
    if anchor.tag == qn("w:sdt"):  # the pandoc --toc block → break so Section 1 is fresh
        toc_break = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
        r.append(br); toc_break.append(r)
        anchor.addnext(toc_break)

    with contextlib.suppress(Exception):  # no running header/footer on the cover
        doc.sections[0].different_first_page_header_footer = True

    doc.save(docx_path)
    return True


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------


def render_markdown_file(
    md_path: str,
    out_dir: str,
    reference_doc: str | None = None,
    *,
    geometry: PageGeometry = DEFAULT_GEOMETRY,
    scale: int = SCALE,
    node_budget: int = NODE_BUDGET,
    toc: bool = True,
    strict: bool = False,
    cover: bool = True,
    cover_image_path: str | None = None,
    logo_path: str | None = None,
) -> RenderResult:
    """Full pipeline for one enriched markdown deliverable → branded .docx.

    The on-disk source markdown is never modified — frontmatter, heading IDs, and the
    opening note are lifted only on the rendered copy (the gate matches by ID OR title)."""
    os.makedirs(out_dir, exist_ok=True)
    stem = Path(md_path).stem
    raw = Path(md_path).read_text(encoding="utf-8")
    warnings: list[str] = []

    meta, body = _split_frontmatter(raw)
    body, title_fallback, note = _transform_headings(body)

    figures_dir = os.path.join(out_dir, "figures")
    attachments_dir = os.path.join(out_dir, "attachments")
    rendered_md, figures = transform(
        body, doc_stem=stem, figures_dir=figures_dir, attachments_dir=attachments_dir,
        geometry=geometry, scale=scale, node_budget=node_budget, strict=strict, warnings=warnings,
    )

    rendered_md_path = os.path.join(out_dir, f"{stem}.rendered.md")
    Path(rendered_md_path).write_text(rendered_md, encoding="utf-8")

    n_tables = n_cards = 0
    docx_path: str | None = os.path.join(out_dir, f"{stem}.docx")
    try:
        to_docx(rendered_md_path, docx_path, reference_doc=reference_doc, toc=toc)
        # Editorial post-process, each guarded so a cosmetic failure never drops the doc.
        # Order matters: style tables BEFORE rendering cards (cards are 1-cell tables too
        # and must not be restyled as data tables).
        try:
            n_tables = style_tables(docx_path)
        except Exception as exc:
            warnings.append(f"table styling skipped: {exc}")
        try:
            n_cards = render_cards(docx_path)
        except Exception as exc:
            warnings.append(f"card rendering skipped: {exc}")
        if cover:
            try:
                add_cover_page(docx_path, meta=meta, title_fallback=title_fallback,
                               cover_image_path=cover_image_path, logo_path=logo_path, note=note)
            except Exception as exc:
                warnings.append(f"cover insertion skipped: {exc}")
    except RenderError as exc:
        if strict:
            raise
        warnings.append(str(exc))
        docx_path = None

    return RenderResult(src=md_path, rendered_md=rendered_md_path, docx=docx_path,
                        figures=figures, warnings=warnings, tables=n_tables, cards=n_cards)


def render_directory(
    src_dir: str,
    out_dir: str,
    reference_doc: str | None = None,
    *,
    pattern: str = "*.md",
    geometry: PageGeometry = DEFAULT_GEOMETRY,
    scale: int = SCALE,
    node_budget: int = NODE_BUDGET,
    toc: bool = True,
    strict: bool = False,
    cover: bool = True,
    cover_image_path: str | None = None,
    logo_path: str | None = None,
) -> list[RenderResult]:
    """Render every ``*.md`` in ``src_dir`` (non-recursive, sorted), skipping our own
    ``*.rendered.md`` intermediates."""
    srcs = sorted(p for p in Path(src_dir).glob(pattern) if not p.name.endswith(".rendered.md"))
    if not srcs:
        raise RenderError(f"no markdown files matching {pattern!r} in {src_dir}")
    results: list[RenderResult] = []
    for src in srcs:
        results.append(render_markdown_file(
            str(src), out_dir, reference_doc=reference_doc, geometry=geometry,
            scale=scale, node_budget=node_budget, toc=toc, strict=strict,
            cover=cover, cover_image_path=cover_image_path, logo_path=logo_path,
        ))
    return results


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m app.render",
        description="Render enriched markdown deliverables (with mermaid/d2 diagrams) "
                    "into branded Word .docx. Diagrams render with LOCAL binaries only "
                    "(mmdc/d2) — source is never sent to a public service.",
    )
    parser.add_argument("deliverables_dir", help="directory of enriched markdown deliverables")
    parser.add_argument("--reference-doc", default=None,
                        help="pandoc branding reference .docx (e.g. assets/branding/techjays-reference.docx)")
    parser.add_argument("--out", required=True, help="output directory for .docx + figures")
    parser.add_argument("--scale", type=int, default=SCALE, help=f"raster scale (default {SCALE})")
    parser.add_argument("--node-budget", type=int, default=NODE_BUDGET,
                        help=f"node count above which a figure is flagged not-readable (default {NODE_BUDGET})")
    parser.add_argument("--no-toc", action="store_true", help="omit the table of contents")
    parser.add_argument("--no-cover", action="store_true", help="skip the designed cover page")
    parser.add_argument("--cover-image", default=None,
                        help="cover hero image (default: techjays-cover.jpg beside --reference-doc)")
    parser.add_argument("--logo", default=None,
                        help="cover logo image (default: techjays-logo.png beside --reference-doc)")
    parser.add_argument("--strict", action="store_true",
                        help="fail if a diagram renderer or pandoc is missing (default: degrade gracefully)")
    args = parser.parse_args(argv)

    status = renderer_status()
    missing = [k for k, v in status.items() if not v]
    if missing:
        print(f"note: local tools not found: {', '.join(missing)} "
              f"({'strict → will fail' if args.strict else 'degrading gracefully'})", file=sys.stderr)

    def _beside_ref(name):
        if args.reference_doc:
            cand = os.path.join(os.path.dirname(os.path.abspath(args.reference_doc)), name)
            return cand if os.path.exists(cand) else None
        return None

    cover_image = args.cover_image or _beside_ref("techjays-cover.jpg")
    logo = args.logo or _beside_ref("techjays-logo.png")

    try:
        results = render_directory(
            args.deliverables_dir, args.out, reference_doc=args.reference_doc,
            scale=args.scale, node_budget=args.node_budget, toc=not args.no_toc, strict=args.strict,
            cover=not args.no_cover, cover_image_path=cover_image, logo_path=logo,
        )
    except RenderError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2

    n_docx = sum(1 for r in results if r.docx)
    n_figs = sum(len(r.figures) for r in results)
    print(f"\nRendered {len(results)} deliverable(s) · {n_docx} .docx · {n_figs} figure(s) → {args.out}")
    for r in results:
        bits = [f"{len(r.figures)} fig"]
        if r.tables:
            bits.append(f"{r.tables} table")
        if r.cards:
            bits.append(f"{r.cards} card")
        nr = sum(f.not_readable for f in r.figures)
        if nr:
            bits.append(f"{nr} view-larger")
        docx = os.path.basename(r.docx) if r.docx else "(no docx — see warnings)"
        print(f"  · {os.path.basename(r.src)} → {docx} · {' · '.join(bits)}")
        for w in r.warnings:
            print(f"      ! {w}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
